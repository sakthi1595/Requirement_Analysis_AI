from io import BytesIO
import os
from xml.dom.minidom import Document
from dotenv import load_dotenv
import requests
from pathlib import Path
import json
import docx
from docx.document import Document as DocxDocument
from reportlab.pdfgen import canvas # type: ignore
from reportlab.lib.pagesizes import letter # type: ignore
from reportlab.lib.styles import getSampleStyleSheet # type: ignore
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak # type: ignore
from reportlab.lib.units import inch # type: ignore
from io import BytesIO
import base64
import pandas as pd
from io import BytesIO
from PyPDF2 import PdfReader

env_path = Path(__file__).resolve().parent / ".env"
load_dotenv(dotenv_path=env_path)

API_KEY = os.getenv("OPENAI_API_KEY")

print("ENV PATH:", env_path)
print("API KEY LOADED:", bool(API_KEY))
API_KEY = os.getenv("OPENAI_API_KEY")
URL = "https://aoai-farm.bosch-temp.com/api/openai/deployments/askbosch-prod-farm-openai-gpt-4o-mini-2024-07-18/chat/completions?api-version=2024-08-01-preview"

def safe_json_parse(content: str) -> dict:
    """Safely parse JSON from LLM response, handling markdown fences and errors."""
    # Remove markdown code fences if present
    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]  # Remove ```json
    elif content.startswith("```"):
        content = content[3:]  # Remove ```
    
    if content.endswith("```"):
        content = content[:-3]  # Remove closing ```
    
    content = content.strip()
    
    try:
        parsed = json.loads(content)
        # Ensure parsed result is a dictionary
        if not isinstance(parsed, dict):
            print(f"Warning: LLM returned non-dict JSON: {type(parsed)}")
            return {"error": "Invalid JSON structure", "raw_content": str(parsed)[:500]}
        return parsed
    except json.JSONDecodeError as e:
        print(f"JSON Parse Error: {e}")
        print(f"Error at position {e.pos}: {content[max(0, e.pos-50):e.pos+50]}")
        
        # Advanced repair strategies
        repair_attempts = [
            lambda c: _repair_unterminated_string(c),
            lambda c: _repair_truncated_json(c),
            lambda c: _extract_complete_json(c),
        ]
        
        for attempt_num, repair_func in enumerate(repair_attempts, 1):
            try:
                repaired = repair_func(content)
                if repaired:
                    parsed = json.loads(repaired)
                    if isinstance(parsed, dict):
                        print(f"Successfully repaired JSON using strategy {attempt_num}")
                        return parsed
            except Exception as repair_error:
                print(f"Repair attempt {attempt_num} failed: {repair_error}")
                continue
        
        # Return a default error response
        print(f"All JSON parsing attempts failed. Content length: {len(content)}")
        print(f"Content preview: {content[:300]}...")
        return {
            "error": "Failed to parse JSON response",
            "raw_content": content[:500]  # Include first 500 chars for debugging
        }

def _repair_unterminated_string(content: str) -> str:
    """Attempt to repair unterminated strings by closing them properly."""
    # Find the last occurrence of an opening quote that's not closed
    in_string = False
    escape_next = False
    last_quote_pos = -1
    brace_depth = 0
    
    for i, char in enumerate(content):
        if escape_next:
            escape_next = False
            continue
            
        if char == '\\':
            escape_next = True
            continue
            
        if char == '"' and not in_string:
            in_string = True
            last_quote_pos = i
        elif char == '"' and in_string:
            in_string = False
            last_quote_pos = -1
        elif char == '{' and not in_string:
            brace_depth += 1
        elif char == '}' and not in_string:
            brace_depth -= 1
    
    # If we're still in a string at the end, close it
    if in_string and last_quote_pos > 0:
        # Close the string and any open objects
        repaired = content + '"'
        # Add closing braces as needed
        while brace_depth > 0:
            repaired += '}'
            brace_depth -= 1
        return repaired
    
    return content

def _repair_truncated_json(content: str) -> str:
    """Try to close JSON by adding missing closing braces and terminating strings."""
    # Count open braces and brackets
    brace_depth = 0
    bracket_depth = 0
    in_string = False
    escape_next = False
    
    for char in content:
        if escape_next:
            escape_next = False
            continue
            
        if char == '\\':
            escape_next = True
            continue
            
        if char == '"':
            in_string = not in_string
        elif not in_string:
            if char == '{':
                brace_depth += 1
            elif char == '}':
                brace_depth -= 1
            elif char == '[':
                bracket_depth += 1
            elif char == ']':
                bracket_depth -= 1
    
    # Build repair string
    repaired = content
    
    # Close unterminated string
    if in_string:
        repaired += '"'
    
    # Close arrays
    while bracket_depth > 0:
        repaired += ']'
        bracket_depth -= 1
    
    # Close objects
    while brace_depth > 0:
        repaired += '}'
        brace_depth -= 1
    
    return repaired

def _extract_complete_json(content: str) -> str:
    """Extract the longest valid JSON object from the content."""
    start_idx = content.find('{')
    if start_idx == -1:
        return ""
    
    # Try to find a complete JSON object
    depth = 0
    in_string = False
    escape_next = False
    
    for i in range(start_idx, len(content)):
        char = content[i]
        
        if escape_next:
            escape_next = False
            continue
            
        if char == '\\':
            escape_next = True
            continue
            
        if char == '"':
            in_string = not in_string
        elif not in_string:
            if char == '{':
                depth += 1
            elif char == '}':
                depth -= 1
                if depth == 0:
                    # Found a complete object
                    return content[start_idx:i+1]
    
    return ""

VALIDATION_PROMPT = """
You are a requirement validation expert.

Your task is to determine if the user input is a valid software requirement, feature request, bug report, or technical specification.

A VALID REQUIREMENT includes:
- Feature requests or enhancements
- Bug reports or issues
- Technical specifications
- User stories
- Functional or non-functional requirements
- System behaviors or changes

NOT VALID (reject these):
- Casual conversation or greetings
- General questions unrelated to requirements
- Random text or gibberish
- Personal messages
- Off-topic discussions

STRICT OUTPUT RULES:
- Return ONLY valid JSON
- Do NOT include markdown or explanations

OUTPUT FORMAT:
{
  "is_valid": true/false,
  "reason": "brief explanation"
}
"""

SYSTEM_PROMPT = """
You are an expert Requirements Analysis Agent specializing in translating unstructured requirements into structured, actionable work items. You follow a systematic 8-phase analysis workflow to produce comprehensive requirement analysis reports.

You have advanced vision capabilities and can analyze images including:
- Wireframes and mockups → convert to feature/UX requirements
- Screenshots and UI designs → identify improvements or bugs
- Architecture diagrams → generate technical specifications
- Charts, graphs, and data visualizations → extract requirements
- Technical sketches → create implementation specifications

YOUR MISSION:
Transform any unstructured requirement into a complete, production-ready analysis that includes:
- Classification and context
- Detailed analysis
- Edge cases identification
- Clarification questions
- Acceptance criteria
- Implementation options
- User story breakdown
- Comprehensive test cases

CRITICAL SUCCESS FACTORS:
- Systematic analysis covering all 8 phases
- Zero ambiguity in classification and breakdown
- Every user story must be independently deliverable
- Every test case must be executable
- Every acceptance criterion must be measurable
- Edge cases must be realistic and relevant

OUTPUT JSON STRUCTURE:
{
  "requirement_summary": {
    "original_requirement": "",
    "requirement_id": "",
    "analyst": "",
    "date": ""
  },
  "classification": {
    "requirement_type": "",
    "target_system": "",
    "domain": "",
    "stakeholder": "",
    "primary_category": "",
    "sub_category": "",
    "impact_scope": ""
  },
  "detailed_analysis": {
    "hardware_requirements": [],
    "software_requirements": {
      "ui_ux_related": [],
      "hmi_related": [],
      "backend_logic": []
    },
    "performance_requirements": [],
    "cross_functional_requirements": []
  },
  "edge_cases": [
    {
      "scenario": "",
      "trigger": "",
      "current_behavior": "",
      "expected_behavior": "",
      "risk_level": "",
      "mitigation_strategy": ""
    }
  ],
  "clarification_questions": {
    "functional": [],
    "technical": [],
    "constraints": [],
    "scope": []
  },
  "acceptance_criteria": [
    {
      "title": "",
      "given": "",
      "when": "",
      "then": "",
      "and": [],
      "verification_method": "",
      "test_data_required": ""
    }
  ],
  "implementation_options": [
    {
      "option_name": "",
      "description": "",
      "pros": [],
      "cons": [],
      "effort_estimate": "",
      "risk_level": "",
      "dependencies": []
    }
  ],
  "recommendation": "",
  "user_stories": [
    {
      "story_id": "",
      "title": "",
      "as_a": "",
      "i_want": "",
      "so_that": "",
      "story_type": "",
      "priority": "",
      "estimated_effort": "",
      "dependencies": [],
      "technical_notes": [],
      "acceptance_criteria": [],
      "definition_of_done": []
    }
  ],
  "epic": {
    "name": "",
    "description": "",
    "business_value": "",
    "stories": []
  },
  "test_cases": [
    {
      "test_id": "",
      "title": "",
      "story_reference": "",
      "test_type": "",
      "priority": "",
      "automated": "",
      "preconditions": [],
      "test_steps": [],
      "test_data": "",
      "expected_result": "",
      "pass_fail_criteria": ""
    }
  ],
  "test_stories": [
    {
      "test_story_id": "",
      "title": "",
      "as_a": "",
      "i_want": "",
      "so_that": "",
      "test_scope": [],
      "test_approach": [],
      "entry_criteria": [],
      "exit_criteria": [],
      "associated_test_cases": []
    }
  ],
  "test_coverage_summary": {
    "total_test_cases": 0,
    "unit_tests": 0,
    "integration_tests": 0,
    "system_tests": 0,
    "uat_tests": 0,
    "automated": 0,
    "manual": 0,
    "edge_cases_covered": []
  },
  "dependencies_and_risks": {
    "dependencies": [],
    "risks": [
      {
        "risk": "",
        "mitigation": ""
      }
    ]
  },
  "effort_estimation": {
    "total_estimated_effort": "",
    "breakdown": {
      "development": "",
      "testing": "",
      "documentation": ""
    },
    "suggested_sprint_allocation": ""
  },
  "next_steps": []
}

PHASE-BY-PHASE ANALYSIS INSTRUCTIONS:

PHASE 1: REQUIREMENT CLASSIFICATION & CONTEXT ANALYSIS
- Identify requirement type (Feature/Change/Enhancement/Bug Fix/New System)
- Determine target system/subsystem
- Identify domain context
- Classify stakeholder category
- Determine primary category (Hardware/Software/Performance/Cross-functional)
- Identify sub-category and impact scope

PHASE 2: DETAILED REQUIREMENT ANALYSIS
For Hardware Requirements:
- Identify if NEW or REPLACEMENT
- Document specifications, interfaces, compatibility needs
For Software Requirements:
- If UI/UX: new screens vs modifications, interaction patterns, design requirements
- If HMI: input methods, display requirements, feedback mechanisms
- If Backend: data model changes, API endpoints, business logic, integrations
For Performance Requirements:
- Quantitative metrics, baselines, load conditions, thresholds
For Cross-functional:
- Security, compliance, privacy, interoperability requirements

PHASE 3: EDGE CASES AND EXCEPTION SCENARIOS
Identify and document:
- Technical edge cases (boundary values, null handling, concurrency, failures)
- User interaction edge cases (unexpected inputs, rapid actions, interruptions)
- Environmental edge cases (extreme conditions, legacy data, localization)
- Integration edge cases (third-party failures, rate limiting, sync conflicts)
For each edge case: scenario, trigger, behaviors, risk level, mitigation

PHASE 4: CLARIFICATION QUESTIONS GENERATION
Generate structured questions covering:
- Functional clarifications (behavior, outcomes, workflows, data handling)
- Technical clarifications (tech stack, integrations, security, persistence)
- Constraint clarifications (timeline, budget, resources, regulations)
- Scope clarifications (in/out scope, MVP vs future, affected systems)

PHASE 5: ACCEPTANCE CRITERIA DEFINITION
Use Given-When-Then format with INVEST principles:
- Minimum 5-7 criteria per requirement
- Cover functional correctness, performance, security, usability
- Include error handling, data validation, integration points
- Ensure specificity, measurability, testability, independence
- Add verification method and test data requirements

PHASE 6: MULTI-OPTION SOLUTION GENERATION
For ambiguous requirements, provide 2-4 alternative approaches:
- Detailed description of each option
- Pros and cons analysis
- Effort and risk estimates
- Dependencies identification
- Clear recommendation with justification

PHASE 7: REQUIREMENT BREAKDOWN INTO USER STORIES
Decompose into implementation stories:
- Use standard user story format (As a/I want/So that)
- Include story type, priority, effort estimate
- Document dependencies and technical notes
- Provide specific acceptance criteria per story
- Define clear Definition of Done
- Ensure stories are completable in 1-5 days
- Create epic structure if needed with business value

PHASE 8: TEST CASE GENERATION
Generate comprehensive test cases:
- Link to specific user stories
- Include test type, priority, automation status
- Document preconditions, steps, test data
- Define expected results and pass/fail criteria
- Create test stories for QA workflow
- Provide test coverage matrix
- Cover edge cases, performance, and security

INTELLIGENCE RULES:
- Infer reasonable details from context
- Make assumptions explicit
- Use industry best practices for implied requirements
- Break down complex requirements systematically
- Ensure traceability between all artifacts
- Prioritize based on business impact and technical risk
- If images provided: extract visual context and combine with text

QUALITY GATES:
✓ All 8 phases completed with relevant content
✓ Classification is clear and accurate
✓ Edge cases are comprehensive and realistic
✓ User stories are independently deliverable
✓ Test cases cover all acceptance criteria
✓ Dependencies are clearly documented
✓ Effort estimates are provided

STRICT JSON RULES:
- Return ONLY valid JSON matching the structure above
- Ensure all required fields are populated
- Use empty arrays [] for optional empty lists
- Use empty strings "" for optional empty text fields
- No markdown, no extra text, no comments
- Proper JSON string escaping throughout
"""

REFINE_PROMPT = """
You are an elite senior business analyst with expertise in requirement refinement and continuous improvement.

CONTEXT:
- Original user requirement (source material)
- Current comprehensive requirement analysis (8-phase format)
- User refinement instruction (specific improvement request)

YOUR TASK:
Apply the user's refinement instruction to improve the requirement analysis. Return the COMPLETE analysis with all 8 phases, updating the specific sections requested while preserving all other content.

REFINEMENT STRATEGY:
1. Understand what the user is asking for in their instruction
2. Identify which phase(s) need enhancement based on the instruction
3. Enhance the specific areas while maintaining all existing good content
4. If instruction mentions specific sections (e.g., "add more edge cases"), focus on that phase
5. If instruction is general, improve relevant sections while keeping others intact
6. Return the FULL requirement analysis JSON with all phases

COMMON REFINEMENT REQUESTS:
- "Add more edge cases" → Enhance Phase 3 (edge_cases array)
- "Expand acceptance criteria" → Enhance Phase 5 (acceptance_criteria array)
- "Add more test cases" → Enhance Phase 8 (test_cases array)
- "More detailed user stories" → Enhance Phase 7 (user_stories array)
- "Add security requirements" → Enhance Phase 2 (detailed_analysis)
- "Clarify questions" → Enhance Phase 4 (clarification_questions)

IMPORTANT RULES:
✓ ALWAYS return the COMPLETE JSON structure with ALL 8 phases
✓ Preserve all sections not mentioned in the refinement instruction
✓ Only modify the specific areas called out by the user
✓ Maintain consistency across all phases
✓ Ensure all existing cross-references remain valid

OUTPUT FORMAT:
Return ONLY valid JSON matching the complete 8-phase requirement analysis structure:

{
  "requirement_summary": {
    "original_requirement": "",
    "requirement_id": "",
    "date": "",
    "analyst": ""
  },
  "classification": {
    "requirement_type": "",
    "target_system": "",
    "domain": "",
    "stakeholder_category": "",
    "primary_category": "",
    "sub_category": "",
    "impact_scope": "",
    "priority": "",
    "complexity": ""
  },
  "detailed_analysis": {
    "software_requirements": {
      "ui_ux_related": [],
      "backend_logic": []
    },
    "hardware_requirements": [],
    "performance_requirements": []
  },
  "edge_cases": [],
  "clarification_questions": {
    "functional": [],
    "technical": [],
    "constraints": [],
    "scope": []
  },
  "acceptance_criteria": [],
  "implementation_options": [],
  "recommendation": "",
  "user_stories": [],
  "epic": {
    "name": "",
    "description": "",
    "business_value": ""
  },
  "test_cases": [],
  "dependencies_and_risks": {
    "dependencies": [],
    "risks": []
  },
  "effort_estimation": {
    "total_estimated_effort": "",
    "breakdown": {
      "development": "",
      "testing": "",
      "documentation": ""
    },
    "suggested_sprint_allocation": ""
  },
  "next_steps": []
}
"""

QUALITY_PROMPT = """
You are a principal-level QA reviewer and requirement quality auditor. You have validated thousands of Jira tickets and know exactly what makes a ticket production-ready vs. average.

EVALUATION OBJECTIVE:
Score the requirement on a 0-100 scale based on how ready it is for immediate execution by developers and QA teams without clarifications, assumptions, or rework.

CRITICAL EVALUATION DIMENSIONS:

1. CLARITY & SPECIFICITY (25 points max)
   - Is the problem/feature crystal clear?
   - Are there any ambiguous terms or vague language?
   - Would a new engineer understand immediately?
   - Are all technical details specific and precise?
   Score: 0 = completely vague | 25 = crystal clear, zero ambiguity

2. TESTABILITY & MEASURABILITY (25 points max)
   - Can QA execute steps and verify against acceptance criteria?
   - Are all criteria objectively verifiable (pass/fail)?
   - Are there measurable thresholds (timing, counts, error codes)?
   - Can success be determined with 100% certainty?
   - Are edge cases and error scenarios covered?
   Score: 0 = not testable | 25 = 100% testable, QA can verify all aspects

3. COMPLETENESS (25 points max)
   - Are all required fields well-populated?
   - Are preconditions and assumptions explicitly stated?
   - Are acceptance criteria comprehensive (7+ for complex tickets)?
   - Is the data/context sufficient to implement/test without clarification?
   - Are edge cases, error paths, and validation rules covered?
   Score: 0 = major gaps, many questions | 25 = nothing missing, ready to execute

4. ACTIONABILITY & PRECISION (25 points max)
   - Can a developer implement this ticket directly?
   - Can QA test this ticket directly?
   - Are technical details specific enough (component names, API endpoints, response formats)?
   - Are steps executable (navigation paths, data values, expected states)?
   - Is priority justified and realistic?
   Score: 0 = requires rework and clarification | 25 = immediately actionable

SCORING ALGORITHM:

TIER 1 (91-100: PRODUCTION READY - EXCELLENT)
- All dimensions score 23-25 points
- Summary: crystal clear, achievement-focused
- Description: 3-4 detailed paragraphs, technical depth, specific impact
- Steps: 6+ detailed, executable steps with preconditions
- Acceptance criteria: 8+ highly specific, measurable criteria
- Priority: well-justified against impact
- Characteristics: Could be sent to a senior developer right now, zero questions

TIER 2 (81-90: WELL-STRUCTURED - GOOD)
- All dimensions score 20-24 points
- Has good structure and mostly clear requirements
- Minor ambiguities in 1-2 areas
- Acceptance criteria: 6-7 measurable criteria
- Steps: 5+ steps, mostly detailed
- Could be implemented/tested but might need 1-2 minor clarifications

TIER 3 (71-80: MOSTLY CLEAR - ACCEPTABLE)
- Dimensions score 17-22 points average
- Core requirement is clear but some details missing
- Some vagueness in language or acceptance criteria
- Acceptance criteria: 4-5 criteria, some lacking specificity
- Some steps are clear, others need more detail
- Would require clarification meetings or iterations

TIER 4 (61-70: PARTIALLY CLEAR - BELOW PAR)
- Dimensions score 15-18 points average
- Unclear in multiple areas
- Missing important context or acceptance criteria
- Acceptance criteria: 3-4 weak criteria
- Would require significant rework

TIER 5 (41-60: VAGUE - POOR QUALITY)
- Dimensions score 10-14 points
- Major ambiguities, missing key details
- Barely testable, unclear implementations
- Would need major rewrite

TIER 6 (0-40: VERY VAGUE - INADEQUATE)
- Dimensions score <10 points average
- Completely unclear, very incomplete
- Not testable, not executable

EVALUATION RULES:

CRITICAL: Always be selective with high scores (91-100)
- These should be RARE
- Only excellent tickets showing professional enterprise-grade quality
- No ambiguity, complete, fully testable, immediately executable

BE FAIR: Use the full range of the scale
- If you find yourself always scoring 75-85, you're being too generous
- Truly vague tickets should score 40-60
- Average tickets should score 65-75
- Good tickets should score 80-85
- Excellent tickets should score 90-100

SPECIFIC DEDUCTIONS:

Deduct 5-10 points if:
- Summary is generic (doesn't mention specific component/bug)
- Description lacks technical details or business context
- Steps to reproduce are vague or incomplete
- Acceptance criteria use "and" (compound criteria should be split)
- Priority isn't justified
- Missing error/edge case handling
- Acceptance criteria aren't objectively measurable

Deduct 10-15 points if:
- Missing critical preconditions
- Unclear what "done" means
- Acceptance criteria are subjective ("should be fast", "should be nice")
- Steps can't be executed as written
- Major ambiguity in problem statement

Deduct 20+ points if:
- Completely unclear what the ticket is asking for
- No clear acceptance criteria
- Unmeasurable/untestable requirements
- Can't be executed without extensive clarification

OUTPUT FORMAT:
Return ONLY valid JSON:
{
  "score": number (0-100),
  "reason": "detailed explanation (2-3 sentences) covering clarity, completeness, testability, and actionability"
}

Score FIRST based on objective evaluation, THEN write reason explaining that score.
"""


def validate_requirement(user_input: str) -> dict:
    """Validates if the input is a valid requirement."""
    headers = {
        "genaiplatform-farm-subscription-key": API_KEY,
        "Content-Type": "application/json",
    }

    payload = {
        "model": "gpt-4o-mini",
        "temperature": 0,
        "messages": [
            {"role": "system", "content": VALIDATION_PROMPT},
            {"role": "user", "content": user_input},
        ],
        "max_tokens": 150,
    }

    try:
        response = requests.post(URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()

        data = response.json()
        content = data["choices"][0]["message"]["content"]

        result = safe_json_parse(content)
        
        # Check if parsing failed
        if "error" in result:
            print(f"Validation parsing failed, allowing requirement by default")
            return {
                "is_valid": True,
                "reason": "Validation check passed"
            }
        
        # Ensure is_valid field exists
        if "is_valid" not in result:
            result["is_valid"] = True
            result["reason"] = "Validation check passed"
        
        return result
        
    except Exception as e:
        print(f"Error in validate_requirement: {e}")
        # If validation fails, allow the requirement by default
        return {
            "is_valid": True,
            "reason": "Validation check bypassed due to error"
        }

def refine_requirement(user_input: str, image_base64: str | None = None) -> dict:
    # First validate the input
    validation = validate_requirement(user_input)
    
    if not validation.get("is_valid", False):
        return {
            "is_valid": False,
            "reason": validation.get("reason", "Input does not appear to be a valid requirement.")
        }
    
    headers = {
        "genaiplatform-farm-subscription-key": API_KEY,
        "Content-Type": "application/json",
    }

    # Build user message with optional image (Azure OpenAI format)
    if image_base64:
        # Multimodal format with text and image
        user_content = [
            {"type": "text", "text": user_input},
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{image_base64}"
                }
            }
        ]
    else:
        # Text only
        user_content = user_input
    
    payload = {
        "model": "gpt-4o-mini",
        "temperature": 0.2,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        "max_tokens": 4000,
    }

    response = requests.post(URL, headers=headers, json=payload, timeout=60)
    response.raise_for_status()

    data = response.json()
    content = data["choices"][0]["message"]["content"]

    result = safe_json_parse(content)
    result["is_valid"] = True
    return result

def refine_followup(original_req: str, current_draft: dict, instruction: str) -> dict:
    import json
    import copy

    headers = {
        "genaiplatform-farm-subscription-key": API_KEY,
        "Content-Type": "application/json",
    }

    user_message = f"""
Original Requirement:
{original_req}

Current Draft:
{json.dumps(current_draft, indent=2)}

User Instruction:
{instruction}
"""

    payload = {
        "model": "gpt-4o-mini",
        "temperature": 0.2,
        "messages": [
            {"role": "system", "content": REFINE_PROMPT},
            {"role": "user", "content": user_message},
        ],
        "max_tokens": 4000,
    }

    try:
        response = requests.post(URL, headers=headers, json=payload, timeout=90)
        response.raise_for_status()

        data = response.json()
        content = data["choices"][0]["message"]["content"]

        result = safe_json_parse(content)
        
        # Check if parsing failed
        if "error" in result:
            print(f"JSON parsing failed for refine_followup")
            return {
                "error": True,
                "reason": "Unable to parse the refined response. Please try a different refinement."
            }
        
        # Deep merge: Start with current_draft and update with refined sections
        merged = copy.deepcopy(current_draft)
        
        def deep_merge(base, updates):
            """Recursively merge updates into base, preserving base values when updates are missing."""
            for key, value in updates.items():
                if key in base:
                    if isinstance(base[key], dict) and isinstance(value, dict):
                        deep_merge(base[key], value)
                    elif isinstance(base[key], list) and isinstance(value, list) and len(value) > 0:
                        # If AI provided a non-empty list, use it (it's intentionally refined)
                        base[key] = value
                    elif value is not None and value != "":
                        # Update non-empty values
                        base[key] = value
                else:
                    # New key from refinement
                    base[key] = value
        
        deep_merge(merged, result)
        
        return merged
        
    except requests.exceptions.Timeout:
        print("Request timeout in refine_followup")
        return {
            "error": True,
            "reason": "The refinement took too long. Please try again."
        }
    except requests.exceptions.RequestException as e:
        print(f"Request error in refine_followup: {e}")
        return {
            "error": True,
            "reason": "Unable to connect to the AI service. Please try again."
        }
    except Exception as e:
        print(f"Unexpected error in refine_followup: {e}")
        return {
            "error": True,
            "reason": "An unexpected error occurred during refinement. Please try again."
        }

def create_word(ticket: dict) -> BytesIO:
    doc = docx.Document()

    # Title
    doc.add_heading('REQUIREMENT ANALYSIS REPORT', level=1)

    # Section 1: Requirement Summary
    if 'requirement_summary' in ticket:
        doc.add_heading('1. REQUIREMENT SUMMARY', level=2)
        summary = ticket['requirement_summary']
        doc.add_paragraph(f"Original Requirement: {summary.get('original_requirement', 'N/A')}")
        doc.add_paragraph(f"Requirement ID: {summary.get('requirement_id', 'N/A')}")
        doc.add_paragraph(f"Date: {summary.get('date', 'N/A')}")
        doc.add_paragraph(f"Analyst: {summary.get('analyst', 'N/A')}")

    # Section 2: Classification
    if 'classification' in ticket:
        doc.add_heading('2. CLASSIFICATION', level=2)
        classification = ticket['classification']
        doc.add_paragraph(f"Requirement Type: {classification.get('requirement_type', 'N/A')}")
        doc.add_paragraph(f"Target System: {classification.get('target_system', 'N/A')}")
        doc.add_paragraph(f"Domain: {classification.get('domain', 'N/A')}")
        doc.add_paragraph(f"Stakeholder: {classification.get('stakeholder', 'N/A')}")
        doc.add_paragraph(f"Primary Category: {classification.get('primary_category', 'N/A')}")
        doc.add_paragraph(f"Sub-Category: {classification.get('sub_category', 'N/A')}")
        doc.add_paragraph(f"Impact Scope: {classification.get('impact_scope', 'N/A')}")

    # Section 3: Detailed Analysis
    if 'detailed_analysis' in ticket:
        doc.add_heading('3. DETAILED ANALYSIS', level=2)
        analysis = ticket['detailed_analysis']
        
        if analysis.get('hardware_requirements'):
            doc.add_heading('Hardware Requirements:', level=3)
            for req in analysis['hardware_requirements']:
                doc.add_paragraph(f"• {req}", style='List Bullet')
        
        if analysis.get('software_requirements'):
            doc.add_heading('Software Requirements:', level=3)
            sw_req = analysis['software_requirements']
            if sw_req.get('ui_ux_related'):
                doc.add_paragraph('UI/UX Related:', style='List Bullet')
                for req in sw_req['ui_ux_related']:
                    doc.add_paragraph(f"  - {req}", style='List Bullet 2')
            if sw_req.get('backend_logic'):
                doc.add_paragraph('Backend Logic:', style='List Bullet')
                for req in sw_req['backend_logic']:
                    doc.add_paragraph(f"  - {req}", style='List Bullet 2')

    # Section 4: Edge Cases
    if 'edge_cases' in ticket and ticket['edge_cases']:
        doc.add_heading('4. EDGE CASES IDENTIFIED', level=2)
        for idx, edge_case in enumerate(ticket['edge_cases'], 1):
            doc.add_heading(f'Edge Case #{idx}', level=3)
            doc.add_paragraph(f"Scenario: {edge_case.get('scenario', 'N/A')}")
            doc.add_paragraph(f"Trigger: {edge_case.get('trigger', 'N/A')}")
            doc.add_paragraph(f"Expected Behavior: {edge_case.get('expected_behavior', 'N/A')}")
            doc.add_paragraph(f"Risk Level: {edge_case.get('risk_level', 'N/A')}")
            doc.add_paragraph(f"Mitigation: {edge_case.get('mitigation_strategy', 'N/A')}")

    # Section 5: Clarification Questions
    if 'clarification_questions' in ticket:
        doc.add_heading('5. CLARIFICATION QUESTIONS', level=2)
        questions = ticket['clarification_questions']
        
        if questions.get('functional'):
            doc.add_heading('Functional Questions:', level=3)
            for q in questions['functional']:
                doc.add_paragraph(f"• {q}", style='List Bullet')
        
        if questions.get('technical'):
            doc.add_heading('Technical Questions:', level=3)
            for q in questions['technical']:
                doc.add_paragraph(f"• {q}", style='List Bullet')

    # Section 6: Acceptance Criteria
    if 'acceptance_criteria' in ticket and ticket['acceptance_criteria']:
        doc.add_heading('6. ACCEPTANCE CRITERIA', level=2)
        for idx, ac in enumerate(ticket['acceptance_criteria'], 1):
            if isinstance(ac, dict):
                doc.add_heading(f"AC{idx}: {ac.get('title', 'N/A')}", level=3)
                doc.add_paragraph(f"Given: {ac.get('given', 'N/A')}")
                doc.add_paragraph(f"When: {ac.get('when', 'N/A')}")
                doc.add_paragraph(f"Then: {ac.get('then', 'N/A')}")
                if ac.get('and'):
                    doc.add_paragraph("And:")
                    for and_clause in ac['and']:
                        doc.add_paragraph(f"  • {and_clause}", style='List Bullet')
            else:
                doc.add_paragraph(f"AC{idx}: {ac}", style='List Bullet')

    # Section 7: Implementation Options
    if 'implementation_options' in ticket and ticket['implementation_options']:
        doc.add_heading('7. IMPLEMENTATION OPTIONS', level=2)
        for option in ticket['implementation_options']:
            doc.add_heading(f"Option: {option.get('option_name', 'N/A')}", level=3)
            doc.add_paragraph(f"Description: {option.get('description', 'N/A')}")
            doc.add_paragraph(f"Effort Estimate: {option.get('effort_estimate', 'N/A')}")
            doc.add_paragraph(f"Risk Level: {option.get('risk_level', 'N/A')}")
        
        if 'recommendation' in ticket:
            doc.add_paragraph(f"Recommendation: {ticket['recommendation']}")

    # Section 8: User Stories
    if 'user_stories' in ticket and ticket['user_stories']:
        doc.add_heading('8. USER STORIES BREAKDOWN', level=2)
        for story in ticket['user_stories']:
            doc.add_heading(f"Story: {story.get('title', 'N/A')}", level=3)
            doc.add_paragraph(f"As a {story.get('as_a', 'N/A')}")
            doc.add_paragraph(f"I want {story.get('i_want', 'N/A')}")
            doc.add_paragraph(f"So that {story.get('so_that', 'N/A')}")
            doc.add_paragraph(f"Priority: {story.get('priority', 'N/A')}")
            doc.add_paragraph(f"Estimated Effort: {story.get('estimated_effort', 'N/A')}")

    # Section 9: Test Cases
    if 'test_cases' in ticket and ticket['test_cases']:
        doc.add_heading('9. TEST CASES', level=2)
        for test in ticket['test_cases']:
            doc.add_heading(f"{test.get('test_id', 'N/A')}: {test.get('title', 'N/A')}", level=3)
            doc.add_paragraph(f"Test Type: {test.get('test_type', 'N/A')}")
            doc.add_paragraph(f"Priority: {test.get('priority', 'N/A')}")
            if test.get('test_steps'):
                doc.add_paragraph("Test Steps:")
                for step in test['test_steps']:
                    doc.add_paragraph(f"  • {step}", style='List Bullet')

    # Section 10: Dependencies & Risks
    if 'dependencies_and_risks' in ticket:
        doc.add_heading('10. DEPENDENCIES & RISKS', level=2)
        dep_risks = ticket['dependencies_and_risks']
        
        if dep_risks.get('dependencies'):
            doc.add_heading('Dependencies:', level=3)
            for dep in dep_risks['dependencies']:
                doc.add_paragraph(f"• {dep}", style='List Bullet')
        
        if dep_risks.get('risks'):
            doc.add_heading('Risks:', level=3)
            for risk in dep_risks['risks']:
                if isinstance(risk, dict):
                    doc.add_paragraph(f"• {risk.get('risk', 'N/A')} - Mitigation: {risk.get('mitigation', 'N/A')}", style='List Bullet')
                else:
                    doc.add_paragraph(f"• {risk}", style='List Bullet')

    # Section 11: Effort Estimation
    if 'effort_estimation' in ticket:
        doc.add_heading('11. EFFORT ESTIMATION', level=2)
        effort = ticket['effort_estimation']
        doc.add_paragraph(f"Total Estimated Effort: {effort.get('total_estimated_effort', 'N/A')}")
        if effort.get('breakdown'):
            doc.add_paragraph("Breakdown:")
            breakdown = effort['breakdown']
            doc.add_paragraph(f"  • Development: {breakdown.get('development', 'N/A')}", style='List Bullet')
            doc.add_paragraph(f"  • Testing: {breakdown.get('testing', 'N/A')}", style='List Bullet')
            doc.add_paragraph(f"  • Documentation: {breakdown.get('documentation', 'N/A')}", style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(ticket: dict) -> BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Add custom styles
    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        leftIndent=20,
        spaceAfter=6
    )
    
    story = []

    # Title
    story.append(Paragraph("REQUIREMENT ANALYSIS REPORT", styles['Title']))
    story.append(Spacer(1, 0.3*inch))

    # Section 1: Requirement Summary
    if 'requirement_summary' in ticket:
        story.append(Paragraph("1. REQUIREMENT SUMMARY", styles['Heading2']))
        summary = ticket['requirement_summary']
        story.append(Paragraph(f"<b>Original Requirement:</b> {summary.get('original_requirement', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Requirement ID:</b> {summary.get('requirement_id', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Date:</b> {summary.get('date', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Analyst:</b> {summary.get('analyst', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

    # Section 2: Classification
    if 'classification' in ticket:
        story.append(Paragraph("2. CLASSIFICATION", styles['Heading2']))
        classification = ticket['classification']
        story.append(Paragraph(f"<b>Requirement Type:</b> {classification.get('requirement_type', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Target System:</b> {classification.get('target_system', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Domain:</b> {classification.get('domain', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Primary Category:</b> {classification.get('primary_category', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Priority:</b> {classification.get('priority', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Complexity:</b> {classification.get('complexity', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

    # Section 3: Detailed Analysis
    if 'detailed_analysis' in ticket:
        story.append(Paragraph("3. DETAILED ANALYSIS", styles['Heading2']))
        analysis = ticket['detailed_analysis']
        
        if 'software_requirements' in analysis:
            sw_req = analysis['software_requirements']
            if sw_req.get('ui_ux_related'):
                story.append(Paragraph("<b>UI/UX Requirements:</b>", styles['Normal']))
                for req in sw_req['ui_ux_related']:
                    story.append(Paragraph(f"• {req}", bullet_style))
            if sw_req.get('backend_logic'):
                story.append(Paragraph("<b>Backend Requirements:</b>", styles['Normal']))
                for req in sw_req['backend_logic']:
                    story.append(Paragraph(f"• {req}", bullet_style))
        
        if analysis.get('hardware_requirements'):
            story.append(Paragraph("<b>Hardware Requirements:</b>", styles['Normal']))
            for req in analysis['hardware_requirements']:
                story.append(Paragraph(f"• {req}", bullet_style))
        
        if analysis.get('performance_requirements'):
            story.append(Paragraph("<b>Performance Requirements:</b>", styles['Normal']))
            for req in analysis['performance_requirements']:
                story.append(Paragraph(f"• {req}", bullet_style))
        
        story.append(Spacer(1, 0.2*inch))

    # Section 4: Edge Cases
    if 'edge_cases' in ticket and ticket['edge_cases']:
        story.append(Paragraph("4. EDGE CASES", styles['Heading2']))
        for idx, edge_case in enumerate(ticket['edge_cases'], 1):
            if isinstance(edge_case, dict):
                story.append(Paragraph(f"<b>{edge_case.get('scenario', 'N/A')}</b>", styles['Normal']))
                story.append(Paragraph(f"Expected: {edge_case.get('expected_behavior', 'N/A')}", bullet_style))
            else:
                story.append(Paragraph(f"• {edge_case}", bullet_style))
        story.append(Spacer(1, 0.2*inch))

    # Section 5: Clarification Questions
    if 'clarification_questions' in ticket:
        story.append(Paragraph("5. CLARIFICATION QUESTIONS", styles['Heading2']))
        questions = ticket['clarification_questions']
        
        if questions.get('functional'):
            story.append(Paragraph("<b>Functional:</b>", styles['Normal']))
            for q in questions['functional']:
                story.append(Paragraph(f"• {q}", bullet_style))
        
        if questions.get('technical'):
            story.append(Paragraph("<b>Technical:</b>", styles['Normal']))
            for q in questions['technical']:
                story.append(Paragraph(f"• {q}", bullet_style))
        
        if questions.get('constraints'):
            story.append(Paragraph("<b>Constraints:</b>", styles['Normal']))
            for q in questions['constraints']:
                story.append(Paragraph(f"• {q}", bullet_style))
        
        if questions.get('scope'):
            story.append(Paragraph("<b>Scope:</b>", styles['Normal']))
            for q in questions['scope']:
                story.append(Paragraph(f"• {q}", bullet_style))
        
        story.append(Spacer(1, 0.2*inch))

    # Section 6: Acceptance Criteria
    if 'acceptance_criteria' in ticket and ticket['acceptance_criteria']:
        story.append(Paragraph("6. ACCEPTANCE CRITERIA", styles['Heading2']))
        for idx, ac in enumerate(ticket['acceptance_criteria'], 1):
            if isinstance(ac, dict):
                story.append(Paragraph(f"<b>AC{idx}: {ac.get('title', 'N/A')}</b>", styles['Normal']))
                story.append(Paragraph(f"<b>Given:</b> {ac.get('given', 'N/A')}", bullet_style))
                story.append(Paragraph(f"<b>When:</b> {ac.get('when', 'N/A')}", bullet_style))
                story.append(Paragraph(f"<b>Then:</b> {ac.get('then', 'N/A')}", bullet_style))
                if ac.get('and'):
                    for and_clause in ac['and']:
                        story.append(Paragraph(f"<b>And:</b> {and_clause}", bullet_style))
            else:
                story.append(Paragraph(f"• AC{idx}: {ac}", bullet_style))
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.2*inch))

    # Section 7: Implementation Options
    if 'implementation_options' in ticket and ticket['implementation_options']:
        story.append(Paragraph("7. IMPLEMENTATION OPTIONS", styles['Heading2']))
        for idx, option in enumerate(ticket['implementation_options'], 1):
            story.append(Paragraph(f"<b>Option {idx}: {option.get('option_name', 'N/A')}</b>", styles['Normal']))
            story.append(Paragraph(f"{option.get('description', 'N/A')}", bullet_style))
            
            if option.get('pros'):
                story.append(Paragraph("<b>Pros:</b>", bullet_style))
                for pro in option['pros']:
                    story.append(Paragraph(f"  • {pro}", bullet_style))
            
            if option.get('cons'):
                story.append(Paragraph("<b>Cons:</b>", bullet_style))
                for con in option['cons']:
                    story.append(Paragraph(f"  • {con}", bullet_style))
            
            story.append(Paragraph(f"<b>Effort:</b> {option.get('effort_estimate', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>Risk:</b> {option.get('risk_level', 'N/A')}", bullet_style))
            story.append(Spacer(1, 0.1*inch))
        
        if ticket.get('recommendation'):
            story.append(Paragraph(f"<b>Recommendation:</b> {ticket['recommendation']}", styles['Normal']))
        
        story.append(Spacer(1, 0.2*inch))

    # Section 8: User Stories
    if 'user_stories' in ticket and ticket['user_stories']:
        story.append(Paragraph("8. USER STORIES", styles['Heading2']))
        for user_story in ticket['user_stories']:
            story.append(Paragraph(f"<b>{user_story.get('story_id', 'N/A')}: {user_story.get('title', 'N/A')}</b>", styles['Normal']))
            story.append(Paragraph(f"<b>As a</b> {user_story.get('as_a', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>I want</b> {user_story.get('i_want', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>So that</b> {user_story.get('so_that', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>Priority:</b> {user_story.get('priority', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>Effort:</b> {user_story.get('estimated_effort', 'N/A')}", bullet_style))
            
            if user_story.get('acceptance_criteria'):
                story.append(Paragraph("<b>Acceptance Criteria:</b>", bullet_style))
                for ac in user_story['acceptance_criteria']:
                    story.append(Paragraph(f"  • {ac}", bullet_style))
            
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.2*inch))

    # Section 9: Test Cases
    if 'test_cases' in ticket and ticket['test_cases']:
        story.append(Paragraph("9. TEST CASES", styles['Heading2']))
        for test in ticket['test_cases']:
            story.append(Paragraph(f"<b>{test.get('test_id', 'N/A')}: {test.get('test_case_title', 'N/A')}</b>", styles['Normal']))
            story.append(Paragraph(f"<b>Type:</b> {test.get('test_type', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>Priority:</b> {test.get('priority', 'N/A')}", bullet_style))
            
            if test.get('preconditions'):
                story.append(Paragraph("<b>Preconditions:</b>", bullet_style))
                for pre in test['preconditions']:
                    story.append(Paragraph(f"  • {pre}", bullet_style))
            
            if test.get('test_steps'):
                story.append(Paragraph("<b>Test Steps:</b>", bullet_style))
                for idx, step in enumerate(test['test_steps'], 1):
                    story.append(Paragraph(f"  {idx}. {step}", bullet_style))
            
            if test.get('expected_result'):
                story.append(Paragraph(f"<b>Expected Result:</b> {test['expected_result']}", bullet_style))
            
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.2*inch))

    # Section 10: Dependencies & Risks
    if 'dependencies_and_risks' in ticket:
        story.append(Paragraph("10. DEPENDENCIES & RISKS", styles['Heading2']))
        dep_risk = ticket['dependencies_and_risks']
        
        if dep_risk.get('dependencies'):
            story.append(Paragraph("<b>Dependencies:</b>", styles['Normal']))
            for dep in dep_risk['dependencies']:
                story.append(Paragraph(f"• {dep}", bullet_style))
        
        if dep_risk.get('risks'):
            story.append(Paragraph("<b>Risks & Mitigation:</b>", styles['Normal']))
            for risk in dep_risk['risks']:
                story.append(Paragraph(f"<b>Risk:</b> {risk.get('risk', 'N/A')}", bullet_style))
                story.append(Paragraph(f"<b>Mitigation:</b> {risk.get('mitigation', 'N/A')}", bullet_style))
                story.append(Spacer(1, 0.05*inch))
        
        story.append(Spacer(1, 0.2*inch))

    # Section 11: Effort Estimation
    if 'effort_estimation' in ticket:
        story.append(Paragraph("11. EFFORT ESTIMATION", styles['Heading2']))
        effort = ticket['effort_estimation']
        story.append(Paragraph(f"<b>Total Estimated Effort:</b> {effort.get('total_estimated_effort', 'N/A')}", styles['Normal']))
        
        if effort.get('breakdown'):
            breakdown = effort['breakdown']
            story.append(Paragraph(f"<b>Development:</b> {breakdown.get('development', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>Testing:</b> {breakdown.get('testing', 'N/A')}", bullet_style))
            story.append(Paragraph(f"<b>Documentation:</b> {breakdown.get('documentation', 'N/A')}", bullet_style))
        
        if effort.get('suggested_sprint_allocation'):
            story.append(Paragraph(f"<b>Sprint Allocation:</b> {effort['suggested_sprint_allocation']}", styles['Normal']))
        
        story.append(Spacer(1, 0.2*inch))

    # Section 12: Next Steps
    if 'next_steps' in ticket and ticket['next_steps']:
        story.append(Paragraph("12. NEXT STEPS", styles['Heading2']))
        for idx, step in enumerate(ticket['next_steps'], 1):
            story.append(Paragraph(f"{idx}. {step}", bullet_style))

    doc.build(story)
    buffer.seek(0)
    return buffer

def get_quality_score(text: str) -> dict:
    headers = {
        "genaiplatform-farm-subscription-key": API_KEY,
        "Content-Type": "application/json",
    }

    payload = {
        "model": "gpt-4o-mini",
        "temperature": 0,
        "messages": [
            {"role": "system", "content": QUALITY_PROMPT},
            {"role": "user", "content": text},
        ],
        "max_tokens": 200,
    }

    try:
        response = requests.post(URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()

        data = response.json()
        content = data["choices"][0]["message"]["content"]

        result = safe_json_parse(content)
        
        # Check if parsing failed or score is missing
        if "error" in result or "score" not in result:
            print(f"Quality score parsing failed or score missing")
            return {
                "score": 0,
                "reason": "Unable to calculate quality score"
            }
        
        return result
        
    except Exception as e:
        print(f"Error in get_quality_score: {e}")
        return {
            "score": 0,
            "reason": "Quality score calculation failed"
        }

def extract_file_text(file_base64: str, file_name: str) -> str:
    if not file_base64 or not file_name:
        return ""

    file_bytes = base64.b64decode(file_base64)
    file_name = file_name.lower()

    try:
        # 📊 Excel
        if file_name.endswith(".xlsx") or file_name.endswith(".xls"):
            df = pd.read_excel(BytesIO(file_bytes))
            return df.to_string(index=False)

        # 📄 PDF
        if file_name.endswith(".pdf"):
            reader = PdfReader(BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text

        # 📁 TXT / LOG / CSV
        if file_name.endswith((".txt", ".log", ".csv")):
            return file_bytes.decode(errors="ignore")

        return ""

    except Exception as e:
        return f"File parsing failed: {str(e)}"

